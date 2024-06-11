# This Python file uses the following encoding: utf-8
# ההערה שכאן מעליי היא קריטית, לא למחוק
import os

from docx import Document

import jinja2
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
from docx.shared import Pt
from docxtpl import DocxTemplate
#by importing these you can make elements and set values quicker
from docx.oxml.shared import OxmlElement, qn

YEAR = 2024  # this is just a dummy variable to showcase that variables can be used as well
file_headers = {}
numb, parasha, date, subj = "","","",""
Font_name, font_size ="", 10

def iter_headings(paragraphs):
    for paragraph in paragraphs:
        name = paragraph.style.name
        if name.startswith('Heading') or name == "Normal" or name == "List Paragraph":
            #print(paragraph.style)
            yield paragraph, paragraph.style.name
        #else:
            #print(paragraph.style.name)


def fix_cs_formatting_runs(run_to_fix, user_cs_font_size, user_cs_font_name, user_is_bold):  #cs: complex script, ex, arabic
        rpr = run_to_fix.element.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rpr.get_or_add_sz()
        szCs = OxmlElement('w:szCs') # size
        sz= OxmlElement('w:sz') # size
        rpr.append(szCs)
        rpr.append(sz)
        lang = OxmlElement('w:lang') #language
        rpr.append(lang)
        if user_is_bold:
            bCs = OxmlElement('w:bCs') #bold the complex language
            rpr.append(bCs)
            bCs.set(qn('w:val'), "True")
            b = OxmlElement('w:b')  # bold the english
            rpr.append(b)
            b.set(qn('w:val'), "True")
        sz.set(qn('w:val'), str(int(user_cs_font_size * 2)))
        szCs.set(qn('w:val'), str(int(user_cs_font_size * 2)))
        lang.set(qn('w:bidi'), 'he-IS')
        rFonts.set(qn('w:cs'), user_cs_font_name)
        rFonts.set(qn('w:ascii'), user_cs_font_name) #you can change the font for the other language
        rFonts.set(qn('w:hAnsi'), user_cs_font_name) #you can change the font for the other language


def do_that(temp_name, file_name, out_name):
    docT = Document(temp_name)
    document = Document(file_name)

    h1, h2, h3 = file_headers[file_name]


    pars = document.paragraphs
    for par in range(len(pars)):
        pars[par].paragraph_format.line_spacing = 1
        if pars[par].text and pars[par].text[0] + pars[par].text[-1] == "()":
            pars[par].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            print(pars[par].text)
        else:
            pars[par].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        pars[par].style = document.styles['Normal']
        runs = pars[par].runs
        for r in range(len(runs)):
            fix_cs_formatting_runs(runs[r], font_size, Font_name, False)
            font = runs[r].font
            font.name = Font_name
            font.cs_size = Pt(font_size)
            font.rtl = True

    document.save("data/c.docx")


    composer = Composer(docT)
    composer.append(document)

    composer.save('data/b.docx')
    tpl = DocxTemplate('data/b.docx')

    context = {
        "סוג": h3,
        "הכותב": h2,
        "שם_מאמר": h1,
        "נושא": subj,

        "מספר": numb,
        "פרשה": parasha,
        "תאריך": date
    }

    jinja_env = jinja2.Environment(autoescape=True)
    tpl.render(context, jinja_env)
    tpl.save(out_name)

    p = "data/imgs/cow.png"
    with open('data/imgs.csv', 'r') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            if row[0] == h3:
                p = "data/imgs/"+row[1]

    try:
        tpl = DocxTemplate('data/b.docx')
        tpl.replace_pic("Replace", p)
        tpl.render({})
        tpl.save(out_name)
    except:
        print ("there no picture to REPLACE. but it's ok")


    docT = Document(out_name)
    new_section = docT.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = docT.sections[-1]

    # Set to 2 column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')
    docT.save(out_name)


def comp(file1, file2, out_name):
    composer = Composer(Document(file1))
    composer.append(Document(file2))
    composer.save(out_name)


def all_files(files):
    for f in files:
        do_that('data/a.docx', f, 'data/output_' + f.split("\\")[-1])
    Document('data/output_' +files[0].split("\\")[-1]).save('last.docx')
    for f in files[1:]:
        comp('last.docx', 'data/output_' + f.split("\\")[-1], 'last.docx')



from tkinter import *
from functools import partial


from easygui import *
from appAlon import *

open_files = fileopenbox("Welcome", filetypes= "*.docx", multiple=True)
#sts = "\\".join(open_files[0].split("\\")[:-1])


root = tk.Tk()

app = MyApp_alon(root, open_files)
root.mainloop()
file_headers = app.file_headers
print(file_headers)

app.pirtey_alon()
root.mainloop()

numb, parasha, date, subj = app.pirtey
print(app.pirtey)


app.Font()
mainloop()

Font_name, font_size = app.font
print(app.font)
# Profile picture

h = max(len(open_files)*37 + 50, 150)
root.geometry("550x" + str(h))  # set starting size of window
#root.maxsize(500, 300)  # width x height
root.config(bg="lightgrey")

buttons = []
Fcli = None
Ncli = None


def clicked(i):
    global Fcli
    if Fcli == None:
        Fcli = i
        buttons[i].config(relief=GROOVE)
    else:
        t = buttons[i]['text']
        buttons[i]['text'] = buttons[Fcli]['text']
        buttons[Fcli]['text'] = t
        buttons[Fcli].config(relief=RAISED)

        t = open_files[i]
        open_files[i] = open_files[Fcli]
        open_files[Fcli] = t

        Fcli = None


def add_l(j, text):
    buttons.append(Button(root, text=text, bg="lightgrey", command=partial(clicked, j-1)))
    buttons[j-1].grid(row=j, column=1, padx=5, pady=5, sticky=E)


# Enter specific information for your profile into the following widgets
enter_info = Label(root, text="תשנה את הסדר בין המאמרים לסדר הרצוי\n שים לב שכדי להחליף סדר עליך ללחוץ על שני מאמרים להחליף את המיקום שלהם", bg="lightgrey")
enter_info.grid(row=0, column=1, columnspan=4, padx=5, pady=5)


files = [file_headers[x][0] + " | " + file_headers[x][1] + " | " + file_headers[x][2] for x in open_files]
for i in range(len(files)):
    add_l(i+1, files[i])

def new_order():
    global files
    print(files)
    files = [x['text'] for x in buttons]
    root.destroy()

end_button = Button(root, text="יאללה לעבודה", command=new_order)
end_button.place(x=450, y=int(h/2))


root.mainloop()


print(file_headers)
all_files(open_files)
#comp('X.docx', 'last.docx', 'last.docx')

doc = Document('data/X.docx')
doc.add_page_break()
composer = Composer(doc)
composer.append(Document('last.docx'))
composer.save('last.docx')



tpl = DocxTemplate('last.docx')


context = {
        "נושא": subj,
        "מספר": numb,
        "פרשה": parasha,
        "תאריך": date
}

jinja_env = jinja2.Environment(autoescape=True)
tpl.render(context, jinja_env)
tpl.save('last.docx')



mypath = "./data"
all_files = os.listdir(mypath)
to_delete = ['b.docx', 'c.docx']
for filename in all_files:
    file_path = os.path.join(mypath, filename)
    if filename in to_delete or "output" in filename:
        os.remove(file_path)

